from __future__ import annotations

from pathlib import Path
from typing import Any, Iterable

from nimo.reporting.renderers.common import money, percent, safe_text
from nimo.reporting.schemas import ReportNarrative


def render_docx(*, evidence: dict[str, Any], narrative: ReportNarrative, output_path: Path) -> Path:
    try:
        from docx import Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("Install Project NIMO with the documents extra to render DOCX reports") from exc

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    palette = {
        "banner": "183E27",
        "accent": "31A247",
        "accent_dark": "1F6D31",
        "surface": "F3F7F4",
        "surface_2": "E7F0E9",
        "border": "D4DED7",
        "ink": RGBColor(22, 35, 27),
        "muted": RGBColor(95, 107, 117),
        "white": RGBColor(255, 255, 255),
    }

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = palette["ink"]
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.15
    for style_name, size, color in [
        ("Title", 24, palette["white"]),
        ("Heading 1", 15, RGBColor(31, 109, 49)),
        ("Heading 2", 11.5, palette["ink"]),
    ]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    period_start = evidence["period"].get("start") or "earliest available"
    period_end = evidence["period"].get("end") or "latest available"
    banner = document.add_table(rows=1, cols=1)
    banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    banner.autofit = False
    banner.columns[0].width = Inches(7.05)
    cell = banner.cell(0, 0)
    _shade_cell(cell, palette["banner"], OxmlElement, qn)
    _set_cell_margins(cell, OxmlElement, qn, top=260, start=360, bottom=260, end=360)
    eyebrow = cell.paragraphs[0]
    eyebrow.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = eyebrow.add_run("PROJECT NIMO")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(210, 237, 217)
    title = cell.add_paragraph("Personal financial report")
    title.style = document.styles["Title"]
    title.paragraph_format.space_before = Pt(5)
    title.paragraph_format.space_after = Pt(5)
    subtitle = cell.add_paragraph(f"{safe_text(evidence['user'])}  |  {period_start} to {period_end}")
    subtitle.paragraph_format.space_after = Pt(0)
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(232, 244, 235)
    document.add_paragraph().paragraph_format.space_after = Pt(0)

    overview = evidence["historical_facts"]["overview"]
    metric_table = document.add_table(rows=2, cols=2)
    metric_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    metric_table.autofit = False
    metric_values = [
        ("TOTAL BALANCE", money(overview.get("total_balance"))),
        ("EXTERNAL INCOME", money(overview.get("total_income"))),
        ("EXTERNAL SPEND", money(overview.get("total_spend"))),
        ("SAVINGS RATE", percent(overview.get("savings_rate"))),
    ]
    for index, (label, value) in enumerate(metric_values):
        row, column = divmod(index, 2)
        cell = metric_table.cell(row, column)
        cell.width = Inches(3.45)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _shade_cell(cell, palette["surface"], OxmlElement, qn)
        _set_cell_border(cell, palette["border"], OxmlElement, qn)
        _set_cell_margins(cell, OxmlElement, qn, top=160, start=220, bottom=160, end=220)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(3)
        label_run = paragraph.add_run(label)
        label_run.font.name = "Arial"
        label_run.font.size = Pt(8)
        label_run.font.bold = True
        label_run.font.color.rgb = palette["muted"]
        value_paragraph = cell.add_paragraph(value)
        value_paragraph.paragraph_format.space_after = Pt(0)
        value_run = value_paragraph.runs[0]
        value_run.font.name = "Arial"
        value_run.font.size = Pt(16)
        value_run.font.bold = True
        value_run.font.color.rgb = RGBColor(31, 109, 49)

    archetype = evidence.get("inferred_behaviour", {}).get("archetype", {})
    callout = document.add_table(rows=1, cols=2)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    callout.autofit = False
    callout.cell(0, 0).width = Inches(1.65)
    callout.cell(0, 1).width = Inches(5.25)
    for item in callout.rows[0].cells:
        _shade_cell(item, palette["surface_2"], OxmlElement, qn)
        _set_cell_border(item, palette["border"], OxmlElement, qn)
        _set_cell_margins(item, OxmlElement, qn, top=120, start=160, bottom=120, end=160)
        item.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    left = callout.cell(0, 0).paragraphs[0]
    left_run = left.add_run("Behaviour profile")
    left_run.font.name = "Arial"
    left_run.font.size = Pt(10)
    left_run.font.bold = True
    left_run.font.color.rgb = RGBColor(31, 109, 49)
    right = callout.cell(0, 1).paragraphs[0]
    right.add_run(
        f"{safe_text(archetype.get('label', 'Insufficient history'))} "
        f"({percent(archetype.get('confidence'), digits=0)} confidence). "
        "This is a descriptive summary, not a judgement or regulated risk classification."
    )

    for title, body in [
        ("Executive summary", narrative.executive_summary),
        ("Financial position", narrative.financial_position),
        ("Cash flow", narrative.cash_flow),
        ("Spending behaviour", narrative.spending_behaviour),
        ("Forecast interpretation", narrative.forecast_interpretation),
        ("Budgets and goals", narrative.budgets_and_goals),
        ("Investing sandbox", narrative.investing),
    ]:
        document.add_heading(title, level=1)
        document.add_paragraph(safe_text(body))

    categories = [
        row for row in evidence["historical_facts"].get("categories", []) if float(row.get("spend") or 0.0) > 0
    ][:10]
    document.add_heading("Leading spending categories", level=1)
    if categories:
        _add_table(
            document,
            ["Category", "Spend", "Transactions"],
            [
                [
                    row.get("category_label", row.get("category", "Uncategorised")),
                    money(row.get("spend")),
                    row.get("transaction_count", row.get("transactions", "")),
                ]
                for row in categories
            ],
            widths=[3.9, 1.55, 1.45],
            palette=palette,
            table_alignment=WD_TABLE_ALIGNMENT,
            vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT,
            paragraph_alignment=WD_ALIGN_PARAGRAPH,
            OxmlElement=OxmlElement,
            qn=qn,
            Inches=Inches,
            Pt=Pt,
        )
    else:
        document.add_paragraph("No category rows were available.")

    forecast = evidence.get("forecast") or {}
    document.add_heading("Forecast summary", level=1)
    if forecast:
        final = forecast.get("final_balance", {})
        _add_table(
            document,
            ["Measure", "Value"],
            [
                ["Scenario", forecast.get("scenario", "Baseline")],
                ["Horizon", f"{forecast.get('horizon_months', '-')} months"],
                ["Median final balance", money(final.get("median"))],
                ["10%-90% final range", f"{money(final.get('p10'))} to {money(final.get('p90'))}"],
                ["Probability below zero", percent(forecast.get("probability_negative_balance"))],
            ],
            widths=[3.4, 3.5],
            palette=palette,
            table_alignment=WD_TABLE_ALIGNMENT,
            vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT,
            paragraph_alignment=WD_ALIGN_PARAGRAPH,
            OxmlElement=OxmlElement,
            qn=qn,
            Inches=Inches,
            Pt=Pt,
        )
    else:
        document.add_paragraph("No stored forecast was included.")

    budget = evidence.get("budget") or {}
    document.add_heading("Budget snapshot", level=1)
    if budget.get("lines"):
        _add_table(
            document,
            ["Budget line", "Budget", "Forecast", "Chance within"],
            [
                [
                    line.get("label", line.get("category_slug", "Budget line")),
                    money(line.get("amount")),
                    money(line.get("forecast_median")),
                    percent(line.get("probability_within_budget")),
                ]
                for line in budget.get("lines", [])
            ],
            widths=[3.15, 1.25, 1.25, 1.25],
            palette=palette,
            table_alignment=WD_TABLE_ALIGNMENT,
            vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT,
            paragraph_alignment=WD_ALIGN_PARAGRAPH,
            OxmlElement=OxmlElement,
            qn=qn,
            Inches=Inches,
            Pt=Pt,
        )
    else:
        document.add_paragraph("No budget evaluation was included.")

    goals = evidence.get("goals") or []
    document.add_heading("Goals", level=1)
    if goals:
        _add_table(
            document,
            ["Goal", "Current", "Target", "Target date", "Probability"],
            [
                [
                    goal.get("name", "Goal"),
                    money(goal.get("current_amount")),
                    money(goal.get("target_amount")),
                    goal.get("target_date", "-"),
                    percent(goal.get("probability_achieved")),
                ]
                for goal in goals
            ],
            widths=[2.25, 1.15, 1.15, 1.25, 1.1],
            palette=palette,
            table_alignment=WD_TABLE_ALIGNMENT,
            vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT,
            paragraph_alignment=WD_ALIGN_PARAGRAPH,
            OxmlElement=OxmlElement,
            qn=qn,
            Inches=Inches,
            Pt=Pt,
        )
    else:
        document.add_paragraph("No active goals were included.")

    investment = evidence.get("investment") or {}
    document.add_heading("Investment simulation", level=1)
    if investment:
        final = investment.get("final_value", {})
        contributed = investment.get("total_contributed", {})
        _add_table(
            document,
            ["Measure", "Value"],
            [
                ["Portfolio", investment.get("portfolio", {}).get("name", "Educational portfolio")],
                ["Horizon", f"{investment.get('horizon_years', '-')} years"],
                ["Median final value", money(final.get("median"))],
                ["10%-90% final range", f"{money(final.get('p10'))} to {money(final.get('p90'))}"],
                ["Median contributions", money(contributed.get("median"))],
                ["Probability below contributions", percent(investment.get("probability_loss_vs_contributions"))],
                ["Probability of negative cash", percent(investment.get("probability_negative_cash"))],
            ],
            widths=[3.4, 3.5],
            palette=palette,
            table_alignment=WD_TABLE_ALIGNMENT,
            vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT,
            paragraph_alignment=WD_ALIGN_PARAGRAPH,
            OxmlElement=OxmlElement,
            qn=qn,
            Inches=Inches,
            Pt=Pt,
        )
        notice = document.add_paragraph()
        notice.paragraph_format.space_before = Pt(5)
        notice.paragraph_format.space_after = Pt(5)
        notice_run = notice.add_run(
            safe_text(
                investment.get(
                    "educational_notice",
                    "Simulated outcomes are not predictions or recommendations.",
                )
            )
        )
        notice_run.italic = True
        notice_run.font.color.rgb = palette["muted"]
    else:
        document.add_paragraph("No investment simulation was included.")

    for title, items in [
        ("Risks", narrative.risks),
        ("Opportunities", narrative.opportunities),
        ("Possible actions", narrative.possible_actions),
        ("Caveats", narrative.caveats),
    ]:
        document.add_heading(title, level=1)
        for item in items or ["None recorded."]:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(safe_text(item))

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(
        f"Project NIMO | Evidence {safe_text(evidence['source_data_version'])} | Educational analysis only"
    )
    footer_run.font.name = "Arial"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = palette["muted"]

    document.core_properties.title = f"Project NIMO financial report - {safe_text(evidence['user'])}"
    document.core_properties.subject = "Educational personal finance analysis"
    document.core_properties.author = "Project NIMO"
    document.save(output_path)
    return output_path


def _add_table(
    document,
    headers: list[str],
    rows: Iterable[Iterable[Any]],
    *,
    widths: list[float],
    palette,
    table_alignment,
    vertical_alignment,
    paragraph_alignment,
    OxmlElement,
    qn,
    Inches,
    Pt,
) -> None:
    materialised = [list(row) for row in rows]
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = table_alignment.CENTER
    table.autofit = False
    _set_repeat_table_header(table.rows[0], OxmlElement, qn)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Inches(widths[index])
        cell.vertical_alignment = vertical_alignment.CENTER
        _shade_cell(cell, palette["surface_2"], OxmlElement, qn)
        _set_cell_border(cell, palette["border"], OxmlElement, qn)
        _set_cell_margins(cell, OxmlElement, qn, top=90, start=120, bottom=90, end=120)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(safe_text(header))
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = palette["ink"]
    for row_index, values in enumerate(materialised):
        cells = table.add_row().cells
        for column, value in enumerate(values):
            cell = cells[column]
            cell.width = Inches(widths[column])
            cell.vertical_alignment = vertical_alignment.CENTER
            if row_index % 2:
                _shade_cell(cell, palette["surface"], OxmlElement, qn)
            _set_cell_border(cell, palette["border"], OxmlElement, qn)
            _set_cell_margins(cell, OxmlElement, qn, top=75, start=120, bottom=75, end=120)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            if column > 0 and _looks_numeric(value):
                paragraph.alignment = paragraph_alignment.RIGHT
            run = paragraph.add_run(safe_text(value))
            run.font.name = "Arial"
            run.font.size = Pt(8.5)
            run.font.color.rgb = palette["ink"]
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _looks_numeric(value: Any) -> bool:
    text = str(value).strip()
    return text.startswith(("£", "$", "€")) or text.endswith("%") or text.replace(",", "").replace(".", "").isdigit()


def _shade_cell(cell, fill: str, OxmlElement, qn) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_border(cell, color: str, OxmlElement, qn) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), color)


def _set_cell_margins(cell, OxmlElement, qn, *, top: int, start: int, bottom: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin_name, margin_value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        tag = f"w:{margin_name}"
        node = margins.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            margins.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row, OxmlElement, qn) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = tr_pr.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        tr_pr.append(marker)
    marker.set(qn("w:val"), "true")
